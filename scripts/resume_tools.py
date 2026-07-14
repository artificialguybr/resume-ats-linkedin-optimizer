#!/usr/bin/env python3
"""resume_tools.py — deterministic helpers for the resume-ats-linkedin-optimizer skill.

Subcommands:
  render     JSON resume  -> single-column, ATS-safe .docx   (needs python-docx)
  validate   resume JSON vs a source (master profile / original) -> truthfulness invariants
  coverage   resume JSON + job-description text -> labeled keyword-coverage ESTIMATE
  deslop     resume JSON -> flag AI-tells / filler / inflated verbs

The model owns the words; this code owns layout, checks, and honest scoring.
validate/coverage/deslop use only the standard library. render needs python-docx.

  python scripts/resume_tools.py <command> --help
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

# ---------------------------------------------------------------------------
# shared helpers
# ---------------------------------------------------------------------------

STANDARD_SECTIONS = ["summary", "experience", "education", "skills", "certifications"]


def load_json(path: str) -> dict:
    try:
        return json.loads(Path(path).read_text(encoding="utf-8"))
    except FileNotFoundError:
        sys.exit(f"error: file not found: {path}")
    except json.JSONDecodeError as e:
        sys.exit(f"error: invalid JSON in {path}: {e}")


def resume_text(resume: dict) -> str:
    """Flatten a resume dict into one lowercase text blob for keyword search."""
    parts: list[str] = []
    b = resume.get("basics", {})
    parts += [b.get("headline", ""), resume.get("summary", "")]
    for exp in resume.get("experience", []):
        parts += [exp.get("title", ""), exp.get("company", "")]
        parts += exp.get("bullets", [])
    for edu in resume.get("education", []):
        parts += [edu.get("degree", ""), edu.get("field", ""), edu.get("institution", "")]
        parts += edu.get("details", [])
    for grp in resume.get("skills", []):
        parts += grp.get("items", [])
    for proj in resume.get("projects", []):
        parts += [proj.get("name", ""), proj.get("description", "")]
        parts += proj.get("bullets", [])
    for cert in resume.get("certifications", []):
        parts += [cert.get("name", ""), cert.get("issuer", "")]
    return "\n".join(p for p in parts if p).lower()


def all_bullets(resume: dict) -> list[str]:
    bullets: list[str] = []
    for exp in resume.get("experience", []):
        bullets += exp.get("bullets", [])
    for proj in resume.get("projects", []):
        bullets += proj.get("bullets", [])
    return bullets


# ---------------------------------------------------------------------------
# validate — truthfulness invariants
# ---------------------------------------------------------------------------

REQUIRED_TOP = ["basics", "experience"]


def norm(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").strip().lower())


def cmd_validate(args) -> int:
    resume = load_json(args.resume)
    failures: list[str] = []
    warnings: list[str] = []

    # 1. shape / schema-lite
    for key in REQUIRED_TOP:
        if key not in resume:
            failures.append(f"missing required top-level section: '{key}'")
    if not resume.get("basics", {}).get("name"):
        failures.append("basics.name is missing")

    if args.source:
        source = load_json(args.source)

        # 2. no fabricated employers
        src_companies = {norm(e.get("company", "")) for e in source.get("experience", [])}
        src_companies |= {norm(e.get("company", "")) for e in source.get("experience", []) if e}
        for exp in resume.get("experience", []):
            c = norm(exp.get("company", ""))
            if c and c not in src_companies:
                failures.append(f"FABRICATED employer not in source: '{exp.get('company')}'")

        # 3. no fabricated titles/dates per matching company
        src_roles = {}
        for e in source.get("experience", []):
            src_roles.setdefault(norm(e.get("company", "")), []).append(e)
        for exp in resume.get("experience", []):
            c = norm(exp.get("company", ""))
            matches = src_roles.get(c, [])
            if matches and norm(exp.get("title", "")) not in {norm(m.get("title", "")) for m in matches}:
                warnings.append(
                    f"title '{exp.get('title')}' at '{exp.get('company')}' not found verbatim in source "
                    "(confirm it is a legitimate rewording, not a promotion inflation)"
                )
            for m in matches:
                if norm(exp.get("start", "")) and norm(m.get("start", "")) and norm(exp.get("start")) != norm(m.get("start")):
                    warnings.append(f"start date differs from source for '{exp.get('company')}' ({exp.get('start')} vs {m.get('start')})")

        # 4. identity untouched
        sb, rb = source.get("basics", {}), resume.get("basics", {})
        for f in ["name", "email", "phone"]:
            if sb.get(f) and rb.get(f) and norm(sb[f]) != norm(rb[f]):
                failures.append(f"identity field '{f}' changed from source ('{sb[f]}' -> '{rb[f]}')")

        # 5. no populated section silently dropped
        for sec in ["experience", "education", "skills", "certifications"]:
            if source.get(sec) and not resume.get(sec):
                warnings.append(f"section '{sec}' populated in source but absent in resume (confirm intentional)")

    # 6. metrics honesty heuristic: flag suspiciously precise unhedged numbers only as info
    # (cannot verify truth here; validate is structural)

    ok = not failures
    report = {
        "invariants_passed": ok,
        "failures": failures,
        "warnings": warnings,
    }
    print(json.dumps(report, indent=2))
    if not ok:
        print("\nFAIL: truthfulness invariants tripped. Reconcile with the user — never invent to fix.", file=sys.stderr)
    return 0 if ok else 1


# ---------------------------------------------------------------------------
# coverage — labeled keyword-coverage estimate
# ---------------------------------------------------------------------------

STOPWORDS = set(
    """a an the and or but if then else for to of in on at by with from into over under
    is are was were be been being have has had do does did will would shall should can
    could may might must this that these those it its as we you your our their they them
    he she his her i me my mine us who whom which what when where why how not no yes all
    any some each every both few more most other such own same so than too very just also
    about above after again against because before below between during down out off up
    while here there once only per via etc""".split()
)

# generic job-posting words that are not real skill keywords
GENERIC_JD = set(
    """experience years work working role roles responsibilities responsible ability able
    strong excellent good great team teams member members company candidate candidates
    join looking seeking required requirement requirements preferred plus bonus nice
    qualifications qualification skills skill knowledge understanding proficiency
    proficient familiarity familiar including include includes etc job position
    opportunity opportunities environment fast paced paced growth grow help support
    ensure ensuring build building develop developing manage managing lead leading
    provide providing deliver delivering across within using use used new best world
    class self starter driven oriented detail passionate motivated communication
    communicate collaborate collaboration stakeholders stakeholder cross functional
    day daily week weekly month monthly year annual etc please apply benefits salary
    location remote hybrid onsite full time part contract senior junior mid level
    domain exposure hands-on handson end""".split()
)

PREFERRED_MARKERS = ("preferred", "nice to have", "nice-to-have", "bonus", "a plus", "plus:", "desirable", "ideally")
REQUIRED_MARKERS = ("required", "requirements", "must have", "must-have", "qualifications", "what you'll need", "you have", "you bring")


def tokenize(text: str) -> list[str]:
    return re.findall(r"[a-zA-Z][a-zA-Z0-9+#.\-]*[a-zA-Z0-9+#]|[a-zA-Z]", text.lower())


def stem(t: str) -> str:
    """Naive plural strip so 'pipelines' matches 'pipeline'. Leaves tech tokens (c++, c#) alone."""
    if len(t) > 4 and t.endswith("ies"):
        return t[:-3] + "y"
    if len(t) > 3 and t.endswith("s") and not t.endswith("ss"):
        return t[:-1]
    return t


def is_content(t: str) -> bool:
    if t in STOPWORDS or t in GENERIC_JD:
        return False
    return len(t) >= 3 or any(c.isdigit() or c in "+#" for c in t)


def extract_keywords(jd_text: str) -> tuple[set[str], set[str]]:
    """Return (required, preferred) keyword sets (unigrams + adjacent bigrams), heuristically tiered.

    Bigrams are formed only from tokens adjacent WITHIN a punctuation-delimited phrase, so
    'Python, SQL' or 'Python and SQL' never produce a spurious 'python sql' keyword.
    """
    tier = "required"  # default
    req: dict[str, int] = {}
    pref: dict[str, int] = {}

    for line in jd_text.splitlines():
        low = line.lower()
        if any(m in low for m in PREFERRED_MARKERS):
            tier = "preferred"
        elif any(m in low for m in REQUIRED_MARKERS):
            tier = "required"
        bucket = pref if tier == "preferred" else req
        # split into phrases on punctuation / conjunctions so bigrams stay meaningful
        for phrase in re.split(r"[,/();:•]|\s[–—]\s|\band\b|\bor\b|\bwith\b", low):
            toks = tokenize(phrase)
            content = [t for t in toks if is_content(t)]
            for t in content:
                bucket[t] = bucket.get(t, 0) + 1
            # bigrams only from directly-adjacent content tokens in the raw phrase
            for a, b in zip(toks, toks[1:]):
                if is_content(a) and is_content(b):
                    bucket[f"{a} {b}"] = bucket.get(f"{a} {b}", 0) + 1

    req_set = {k for k in req}
    pref_set = {k for k in pref} - req_set
    return req_set, pref_set


def coverage_of(keywords: set[str], resume_txt: str) -> tuple[set[str], set[str]]:
    """A unigram matches on stemmed token presence; a bigram matches if the phrase appears
    OR both stemmed tokens are present anywhere (candidate clearly has both skills)."""
    rtokens = {stem(t) for t in tokenize(resume_txt)}
    present, missing = set(), set()
    for kw in keywords:
        if " " in kw:
            a, b = kw.split(" ", 1)
            hit = kw in resume_txt or (stem(a) in rtokens and stem(b) in rtokens)
        else:
            hit = stem(kw) in rtokens
        (present if hit else missing).add(kw)
    return present, missing


def cmd_coverage(args) -> int:
    resume = load_json(args.resume)
    jd_text = Path(args.jd).read_text(encoding="utf-8") if Path(args.jd).exists() else args.jd
    rtext = resume_text(resume)

    req, pref = extract_keywords(jd_text)
    req_present, req_missing = coverage_of(req, rtext)
    pref_present, pref_missing = coverage_of(pref, rtext)

    req_cov = len(req_present) / len(req) if req else 1.0
    pref_cov = len(pref_present) / len(pref) if pref else 1.0

    bullets = all_bullets(resume)
    quant = sum(1 for b in bullets if re.search(r"\d", b)) / len(bullets) if bullets else 0.0

    present_sections = sum(1 for s in STANDARD_SECTIONS if resume.get(s))
    section_completeness = present_sections / len(STANDARD_SECTIONS)

    # keyword distribution: fraction of content sections that contain >=1 JD keyword
    all_kw = req | pref
    content_sections = {
        "summary": resume.get("summary", ""),
        "experience": " ".join(all_bullets(resume)),
        "skills": " ".join(i for g in resume.get("skills", []) for i in g.get("items", [])),
    }
    content_sections = {k: v for k, v in content_sections.items() if v}
    with_kw = sum(1 for v in content_sections.values() if any(kw in v.lower() for kw in all_kw))
    distribution = with_kw / len(content_sections) if content_sections else 0.0

    score = (
        0.40 * req_cov
        + 0.20 * pref_cov
        + 0.20 * quant
        + 0.10 * section_completeness
        + 0.10 * distribution
    )

    band = "strong" if score >= 0.85 else "moderate" if score >= 0.70 else "weak"

    report = {
        "_note": "ESTIMATE of keyword fit + quality — NOT a real ATS internal score.",
        "estimated_score_pct": round(score * 100, 1),
        "band": band,
        "components": {
            "required_coverage_pct": round(req_cov * 100, 1),
            "preferred_coverage_pct": round(pref_cov * 100, 1),
            "quantification_rate_pct": round(quant * 100, 1),
            "section_completeness_pct": round(section_completeness * 100, 1),
            "keyword_distribution_pct": round(distribution * 100, 1),
        },
        "missing_required_keywords": sorted(req_missing),
        "missing_preferred_keywords": sorted(pref_missing),
        "advice": "Add missing REQUIRED keywords you genuinely qualify for (in Skills + a real bullet). Do not stuff or fabricate.",
    }
    print(json.dumps(report, indent=2))
    return 0


# ---------------------------------------------------------------------------
# deslop — flag AI-tells / filler / inflated verbs
# ---------------------------------------------------------------------------

INFLATED = {
    "spearheaded": "led", "orchestrated": "coordinated", "architected": "designed",
    "leveraged": "used", "utilized": "used", "facilitated": "ran/helped",
    "championed": "advocated for", "pioneered": "introduced", "harnessed": "used",
    "helmed": "led", "navigated": "managed",
}
BUZZWORDS = {
    "synergy", "synergize", "paradigm", "robust", "scalable", "actionable", "impactful",
    "best-in-class", "cutting-edge", "world-class", "next-level", "results-driven",
    "detail-oriented", "team player", "hard worker", "go-getter", "self-starter",
    "dynamic", "passionate", "holistic", "seamless", "game-changer", "value-add",
    "mission-critical", "deep dive", "move the needle", "low-hanging fruit",
    "think outside the box",
}
FILLER = {
    "in order to": "to", "on a daily basis": "daily", "due to the fact that": "because",
    "a wide variety of": "many", "responsible for": "(lead with an action verb)",
    "tasked with": "(lead with an action verb)", "successfully": "(delete — show it)",
    "various": "(name them or delete)", "helped": "(be specific)", "worked on": "(be specific)",
}


def cmd_deslop(args) -> int:
    resume = load_json(args.resume)
    findings: list[dict] = []
    texts = [("summary", resume.get("summary", ""))]
    for i, exp in enumerate(resume.get("experience", [])):
        for j, b in enumerate(exp.get("bullets", [])):
            texts.append((f"experience[{i}].bullet[{j}]", b))
    for i, proj in enumerate(resume.get("projects", [])):
        for j, b in enumerate(proj.get("bullets", [])):
            texts.append((f"projects[{i}].bullet[{j}]", b))

    for loc, text in texts:
        low = text.lower()
        for word, repl in INFLATED.items():
            if re.search(rf"\b{re.escape(word)}\b", low):
                findings.append({"location": loc, "type": "inflated-verb", "found": word, "suggest": repl})
        for bw in BUZZWORDS:
            # word-boundary match for single words (so 'dynamic' != 'dynamically'); substring for phrases
            hit = (" " in bw or "-" in bw) and bw in low or (" " not in bw and "-" not in bw and re.search(rf"\b{re.escape(bw)}\b", low))
            if hit:
                findings.append({"location": loc, "type": "buzzword", "found": bw, "suggest": "cut / replace with concrete proof"})
        for f, repl in FILLER.items():
            if f in low:
                findings.append({"location": loc, "type": "filler", "found": f, "suggest": repl})
        if "—" in text or "--" in text:
            findings.append({"location": loc, "type": "punctuation", "found": "em-dash/--", "suggest": "use comma/period/'to'"})

    report = {"slop_findings": findings, "clean": not findings}
    print(json.dumps(report, indent=2))
    return 0


# ---------------------------------------------------------------------------
# render — JSON -> ATS-safe .docx
# ---------------------------------------------------------------------------

def cmd_render(args) -> int:
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
    except ImportError:
        sys.exit("error: python-docx is required for render. Install:  pip install python-docx")

    resume = load_json(args.resume)
    b = resume.get("basics", {})

    doc = Document()
    # single column, 1-inch margins, web-safe font
    for section in doc.sections:
        section.left_margin = section.right_margin = Inches(1.0)
        section.top_margin = section.bottom_margin = Inches(0.9)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    # ensure east-asian font mapping doesn't override
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), "Calibri")
    rfonts.set(qn("w:hAnsi"), "Calibri")

    def heading(text: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(11.5)
        # bottom border
        ppr = p._p.get_or_add_pPr()
        pbdr = ppr.makeelement(qn("w:pBdr"), {})
        bottom = pbdr.makeelement(qn("w:bottom"), {
            qn("w:val"): "single", qn("w:sz"): "6", qn("w:space"): "1", qn("w:color"): "999999",
        })
        pbdr.append(bottom)
        ppr.append(pbdr)

    def bullet(text: str):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(text)

    # --- header: name + contact in BODY (never a header/footer) ---
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(0)
    nr = name_p.add_run(b.get("name", ""))
    nr.bold = True
    nr.font.size = Pt(18)
    if b.get("headline"):
        hp = doc.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hp.paragraph_format.space_after = Pt(0)
        hr = hp.add_run(b["headline"])
        hr.font.size = Pt(11.5)

    contact_bits = [b.get(k) for k in ["phone", "email", "location", "linkedin", "github", "website"] if b.get(k)]
    if contact_bits:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(4)
        cp.add_run("  |  ".join(contact_bits)).font.size = Pt(9.5)

    # --- summary ---
    if resume.get("summary"):
        heading("Summary")
        doc.add_paragraph(resume["summary"])

    # --- experience ---
    if resume.get("experience"):
        heading("Work Experience")
        for exp in resume["experience"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(f"{exp.get('title','')} — {exp.get('company','')}")
            r.bold = True
            meta = " | ".join(x for x in [exp.get("location", ""), f"{exp.get('start','')} – {exp.get('end','')}".strip(" –")] if x)
            if meta:
                mp = doc.add_paragraph()
                mp.paragraph_format.space_after = Pt(2)
                mr = mp.add_run(meta)
                mr.italic = True
                mr.font.size = Pt(9.5)
            for bl in exp.get("bullets", []):
                bullet(bl)

    # --- projects ---
    if resume.get("projects"):
        heading("Projects")
        for proj in resume["projects"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            title = proj.get("name", "")
            if proj.get("link"):
                title += f" ({proj['link']})"
            p.add_run(title).bold = True
            if proj.get("description"):
                doc.add_paragraph(proj["description"])
            for bl in proj.get("bullets", []):
                bullet(bl)

    # --- skills ---
    if resume.get("skills"):
        heading("Skills")
        for grp in resume["skills"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.add_run(f"{grp.get('category','')}: ").bold = True
            p.add_run(", ".join(grp.get("items", [])))

    # --- education ---
    if resume.get("education"):
        heading("Education")
        for edu in resume["education"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            deg = ", ".join(x for x in [edu.get("degree", ""), edu.get("field", "")] if x)
            p.add_run(f"{deg} — {edu.get('institution','')}".strip(" —")).bold = True
            tail = " | ".join(x for x in [edu.get("location", ""), f"{edu.get('start','')} – {edu.get('end','')}".strip(" –")] if x)
            if tail:
                tp = doc.add_paragraph()
                tp.paragraph_format.space_after = Pt(2)
                tr = tp.add_run(tail)
                tr.italic = True
                tr.font.size = Pt(9.5)
            for d in edu.get("details", []):
                bullet(d)

    # --- certifications ---
    if resume.get("certifications"):
        heading("Certifications")
        for c in resume["certifications"]:
            line = c.get("name", "")
            extra = ", ".join(x for x in [c.get("issuer", ""), c.get("year", "")] if x)
            if extra:
                line += f" — {extra}"
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            p.add_run(line)

    # --- awards ---
    if resume.get("awards"):
        heading("Awards")
        for a in resume["awards"]:
            line = a.get("name", "")
            extra = ", ".join(x for x in [a.get("issuer", ""), a.get("year", "")] if x)
            if extra:
                line += f" — {extra}"
            doc.add_paragraph(line)

    out = args.out or "resume.docx"
    doc.save(out)
    print(f"wrote {out}  (single-column, ATS-safe). Export to PDF for submission; keep .docx for portals that request it.")
    return 0


# ---------------------------------------------------------------------------
# cli
# ---------------------------------------------------------------------------

def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    sub = parser.add_subparsers(dest="command", required=True)

    p_render = sub.add_parser("render", help="JSON resume -> ATS-safe .docx (needs python-docx)")
    p_render.add_argument("--resume", required=True)
    p_render.add_argument("--out", default=None)
    p_render.set_defaults(func=cmd_render)

    p_val = sub.add_parser("validate", help="truthfulness invariants (resume vs source)")
    p_val.add_argument("--resume", required=True)
    p_val.add_argument("--source", default=None, help="master-profile.json or original resume.json to check against")
    p_val.set_defaults(func=cmd_validate)

    p_cov = sub.add_parser("coverage", help="labeled keyword-coverage ESTIMATE vs a job description")
    p_cov.add_argument("--resume", required=True)
    p_cov.add_argument("--jd", required=True, help="path to a job-description .txt (or the text itself)")
    p_cov.set_defaults(func=cmd_coverage)

    p_ds = sub.add_parser("deslop", help="flag AI-tells / filler / inflated verbs")
    p_ds.add_argument("--resume", required=True)
    p_ds.set_defaults(func=cmd_deslop)

    args = parser.parse_args()
    return args.func(args)


if __name__ == "__main__":
    raise SystemExit(main())
